import os
import uuid
import structlog
from typing import List, Optional
from datetime import datetime
from fastapi import UploadFile
from sqlmodel import Session, select
import chromadb
from chromadb.config import Settings
from pypdf import PdfReader
from docx import Document as DocxDocument

from app.models import Document, DocumentType
from app.database import engine

logger = structlog.get_logger(__name__)

class DocumentService:
    def __init__(self):
        # Initialize ChromaDB
        self.chroma_client = chromadb.PersistentClient(path="./amplified_vectors")
        self.collection = self.chroma_client.get_or_create_collection(name="documents")
        
    async def process_upload(self, file: UploadFile, type: str, tags: List[str], user_id: str) -> Document:
        """
        Process uploaded file: save to DB, extract text, generate embeddings
        """
        content = await file.read()
        filename = file.filename
        
        # 1. Extract Text
        text_content = self._extract_text(content, filename)
        
        if not text_content:
            raise ValueError("Could not extract text from file")
            
        # 2. Save to SQL DB
        doc_id = str(uuid.uuid4())
        document = Document(
            id=doc_id,
            user_id=user_id,
            name=filename,
            type=type,
            extracted_text=text_content,
            tags=",".join(tags) if tags else "",
            created_at=datetime.now()
        )
        
        with Session(engine) as session:
            session.add(document)
            session.commit()
            session.refresh(document)
            
        # 3. Generate Embeddings & Store in Vector DB
        # Split text into chunks (simple splitting for now)
        chunks = self._chunk_text(text_content)
        
        ids = [f"{doc_id}_{i}" for i in range(len(chunks))]
        metadatas = [{"doc_id": doc_id, "user_id": user_id, "filename": filename, "type": type} for _ in chunks]
        
        self.collection.add(
            documents=chunks,
            metadatas=metadatas,
            ids=ids
        )
        
        logger.info(f"Processed document {filename} with {len(chunks)} chunks")
        return document

    def _extract_text(self, content: bytes, filename: str) -> str:
        """Extract text based on file extension"""
        import io
        
        if filename.lower().endswith('.pdf'):
            try:
                reader = PdfReader(io.BytesIO(content))
                text = ""
                for page in reader.pages:
                    text += page.extract_text() + "\n"
                return text
            except Exception as e:
                logger.error(f"PDF extraction error: {e}")
                return ""
                
        elif filename.lower().endswith('.docx'):
            try:
                doc = DocxDocument(io.BytesIO(content))
                text = "\n".join([para.text for para in doc.paragraphs])
                return text
            except Exception as e:
                logger.error(f"DOCX extraction error: {e}")
                return ""
                
        else:
            # Assume text/plain
            try:
                return content.decode('utf-8', errors='ignore')
            except Exception:
                return ""

    def _chunk_text(self, text: str, chunk_size: int = 1000, overlap: int = 100) -> List[str]:
        """Simple text chunking"""
        chunks = []
        start = 0
        while start < len(text):
            end = start + chunk_size
            chunks.append(text[start:end])
            start = end - overlap
        return chunks

    def search_documents(self, query: str, limit: int = 5, user_id: str = None) -> List[dict]:
        """Search for relevant document chunks"""
        # Build where filter for user_id
        where_filter = {"user_id": user_id} if user_id else None
        
        results = self.collection.query(
            query_texts=[query],
            n_results=limit,
            where=where_filter
        )
        
        # Format results
        formatted_results = []
        if results['documents']:
            for i, doc_text in enumerate(results['documents'][0]):
                metadata = results['metadatas'][0][i]
                formatted_results.append({
                    "document_id": metadata['doc_id'],
                    "filename": metadata['filename'],
                    "snippet": doc_text,
                    "score": 0.0 # Chroma doesn't return cosine score by default in this simple API usage without distance
                })
                
        return formatted_results

    def list_documents(
        self, 
        user_id: str, 
        type: Optional[str] = None,
        limit: int = 50,
        offset: int = 0
    ) -> List[Document]:
        with Session(engine) as session:
            statement = select(Document).where(
                Document.user_id == user_id
            ).order_by(Document.created_at.desc()).offset(offset).limit(limit)
            if type:
                statement = statement.where(Document.type == type)
            results = session.exec(statement).all()
            return results

    def delete_document(self, document_id: str, user_id: str):
        with Session(engine) as session:
            # Verify ownership
            statement = select(Document).where(
                Document.id == document_id,
                Document.user_id == user_id
            )
            doc = session.exec(statement).first()
            if doc:
                session.delete(doc)
                session.commit()
                
        # Delete from Vector DB
        # Note: Chroma delete by metadata is supported
        self.collection.delete(
            where={"doc_id": document_id, "user_id": user_id}
        )
